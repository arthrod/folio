"""Generate three small, snappy, thematically-distinct .docx pairs for the
redline3 playground "play with an example" cards.

Each pair (A + B) shares structure, but B is a realistic revision: ~20% of
paragraphs reworded, table values/descriptions changed (sometimes a row added),
one figure swapped, and the odd paragraph inserted or deleted. Images are drawn
locally with Pillow (deterministic gradients with a figure label) so the script
never depends on the network. Small on purpose (~6-10 pages) so the compare and
render feel instant.

Usage:  python gen-redline3-quick.py <out_dir>
"""

from __future__ import annotations

import io
import random
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

OUT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(".")
OUT_DIR.mkdir(parents=True, exist_ok=True)

# --- Word pools give each theme its own texture. --------------------------------

BASE_WORDS = (
    "agreement clause party liability warranty covenant termination confidential "
    "obligation governing remedy breach provision counterparty schedule appendix "
    "representation disclosure amendment consideration enforceable severable "
    "assignment notice waiver compliance instrument delivery acceptance milestone "
    "deliverable precedent statutory equitable fiduciary resolution ratify rescind "
    "encumbrance collateral guarantor recital whereas hereto therein hereunder"
).split()

THEMES = {
    "services": {
        "title": "Master Services Agreement",
        "accent": (37, 99, 235),
        "headings": [
            "Definitions and Interpretation", "Scope of Services", "Service Levels",
            "Fees and Payment", "Deliverables and Milestones", "Warranties",
            "Limitation of Liability", "Term and Termination", "Governing Law",
        ],
        "extra": "engagement statement-of-work acceptance escalation onboarding uptime".split(),
    },
    "lease": {
        "title": "Commercial Lease Agreement",
        "accent": (217, 119, 6),
        "headings": [
            "Demised Premises", "Term of Lease", "Rent and Escalation",
            "Use and Occupancy", "Maintenance and Repairs", "Insurance",
            "Assignment and Subletting", "Default and Remedies", "Surrender",
        ],
        "extra": "premises landlord tenant rent leasehold fixtures easement zoning".split(),
    },
    "nda": {
        "title": "Mutual Non-Disclosure Agreement",
        "accent": (124, 58, 237),
        "headings": [
            "Purpose", "Definition of Confidential Information", "Obligations",
            "Permitted Disclosures", "Term of Confidentiality", "Return of Materials",
            "No License Granted", "Injunctive Relief", "Miscellaneous",
        ],
        "extra": "confidential disclosure recipient discloser trade-secret proprietary".split(),
    },
}


def sentence(rng: random.Random, words: list[str], n: int) -> str:
    ws = [rng.choice(words) for _ in range(n)]
    ws[0] = ws[0].capitalize()
    return " ".join(ws) + "."


def paragraph(rng: random.Random, words: list[str]) -> str:
    return " ".join(sentence(rng, words, rng.randint(9, 18)) for _ in range(rng.randint(3, 5)))


def make_image(accent: tuple[int, int, int], label: str, n: int) -> bytes:
    """A deterministic gradient tile with a figure label — no network needed."""
    w, h = 880, 520
    img = Image.new("RGB", (w, h))
    px = img.load()
    ar, ag, ab = accent
    for y in range(h):
        t = y / h
        for x in range(w):
            s = (x / w) * 0.5 + t * 0.5
            px[x, y] = (
                int(ar * (1 - s) + 245 * s),
                int(ag * (1 - s) + 245 * s),
                int(ab * (1 - s) + 245 * s),
            )
    draw = ImageDraw.Draw(img)
    # A few translucent bands for texture.
    for i in range(5):
        yy = int(h * (0.15 + i * 0.17))
        draw.line([(0, yy), (w, yy - 60)], fill=(255, 255, 255), width=2)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 68)
    except Exception:  # noqa: BLE001
        font = ImageFont.load_default()
    draw.text((44, h - 110), label, fill=(255, 255, 255), font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def build_model(theme_key: str, theme: dict, images: list[bytes]) -> list[dict]:
    rng = random.Random(hash(theme_key) & 0xFFFF)
    words = BASE_WORDS + theme["extra"]
    blocks: list[dict] = [{"type": "title", "text": theme["title"]}]
    for s in range(6):
        blocks.append({"type": "heading", "text": f"{s + 1}. {theme['headings'][s % len(theme['headings'])]}"})
        for _ in range(rng.randint(3, 5)):
            blocks.append({"type": "para", "text": paragraph(rng, words)})
        if s % 2 == 0:
            rows = [["Ref", "Description", "Owner", "Value"]]
            for r in range(rng.randint(3, 5)):
                rows.append([
                    f"{s + 1}.{r + 1}",
                    sentence(rng, words, rng.randint(3, 6)),
                    rng.choice(words).capitalize(),
                    f"${rng.randint(1, 950)}k",
                ])
            blocks.append({"type": "table", "rows": rows})
        if s % 3 == 1 and images:
            blocks.append({"type": "image", "idx": s % len(images),
                           "caption": f"Figure {s + 1}: {sentence(rng, words, 4)}"})
    return blocks


def revise_model(theme_key: str, blocks: list[dict], images: list[bytes]) -> list[dict]:
    rng = random.Random((hash(theme_key) & 0xFFFF) ^ 0x5A5A)
    words = BASE_WORDS + THEMES[theme_key]["extra"]
    out: list[dict] = []
    for b in blocks:
        nb = {**b}
        if b["type"] == "para" and rng.random() < 0.20:
            nb["text"] = paragraph(rng, words)
        elif b["type"] == "table":
            rows = [list(b["rows"][0])]
            for row in b["rows"][1:]:
                r = list(row)
                if rng.random() < 0.55:
                    r[3] = f"${rng.randint(1, 950)}k"
                if rng.random() < 0.35:
                    r[1] = sentence(rng, words, rng.randint(3, 6))
                rows.append(r)
            if rng.random() < 0.5:
                rows.append([f"{rows[-1][0]}x", sentence(rng, words, 4),
                             rng.choice(words).capitalize(), f"${rng.randint(1, 950)}k"])
            nb["rows"] = rows
        elif b["type"] == "image" and images and rng.random() < 0.6:
            nb["idx"] = (b["idx"] + 2) % len(images)
        if b["type"] == "para" and rng.random() < 0.06:
            continue  # deletion
        out.append(nb)
        if b["type"] == "para" and rng.random() < 0.06:
            out.append({"type": "para", "text": paragraph(rng, words)})  # insertion
    return out


def render(blocks: list[dict], images: list[bytes], path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(11)
    for b in blocks:
        if b["type"] == "title":
            h = doc.add_heading(b["text"], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif b["type"] == "heading":
            doc.add_heading(b["text"], level=1)
        elif b["type"] == "para":
            doc.add_paragraph(b["text"])
        elif b["type"] == "table":
            rows = b["rows"]
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    t.rows[ri].cells[ci].text = str(cell)
            doc.add_paragraph()
        elif b["type"] == "image":
            doc.add_picture(io.BytesIO(images[b["idx"]]), width=Inches(4.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(b.get("caption", ""))
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x73, 0x73, 0x73)
    doc.save(path)
    print(f"wrote {path.name} ({path.stat().st_size // 1024} KB, {len(blocks)} blocks)", file=sys.stderr)


for i, (key, theme) in enumerate(THEMES.items(), start=1):
    imgs = [make_image(theme["accent"], f"{theme['title']} · Fig {j + 1}", j) for j in range(4)]
    model_a = build_model(key, theme, imgs)
    model_b = revise_model(key, model_a, imgs)
    render(model_a, imgs, OUT_DIR / f"pair{i}-a.docx")
    render(model_b, imgs, OUT_DIR / f"pair{i}-b.docx")
print("done", file=sys.stderr)
