"""Generate two large (~200pp) .docx sample documents for the redline3 playground.

A and B share structure but B is a realistic revision: ~18% of paragraphs
reworded, table values changed, some images swapped, a few blocks inserted/
deleted. Random images are downloaded from picsum.photos. Deterministic (seeded)
so re-runs are reproducible.
"""

from __future__ import annotations

import io
import random
import sys
import urllib.request
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(sys.argv[1])
SECTIONS = int(sys.argv[2]) if len(sys.argv) > 2 else 215
SUFFIX = sys.argv[3] if len(sys.argv) > 3 else ""
OUT_DIR.mkdir(parents=True, exist_ok=True)

RNG = random.Random(20260717)

WORDS = (
    "agreement clause party liability indemnity warranty covenant jurisdiction "
    "termination confidential obligation governing arbitration remedy breach "
    "provision counterparty schedule appendix representation disclosure diligence "
    "amendment consideration enforceable severable assignment notice waiver "
    "compliance regulatory framework instrument perpetual royalty license premises "
    "delivery acceptance milestone deliverable escalation mitigation stipulation "
    "precedent statutory equitable fiduciary quorum resolution ratify rescind "
    "encumbrance collateral guarantor solvent recital whereas hereto therein"
).split()

HEADINGS = [
    "Definitions and Interpretation", "Scope of Engagement", "Representations",
    "Warranties", "Indemnification", "Limitation of Liability", "Confidentiality",
    "Intellectual Property", "Term and Termination", "Governing Law",
    "Dispute Resolution", "Data Protection", "Force Majeure", "Assignment",
    "Notices", "Payment Terms", "Milestones and Deliverables", "Audit Rights",
    "Compliance", "Schedules and Appendices",
]

IMAGE_SEEDS = [f"folio-{i}" for i in range(40)]


def sentence(rng: random.Random, n_words: int) -> str:
    ws = [rng.choice(WORDS) for _ in range(n_words)]
    ws[0] = ws[0].capitalize()
    return " ".join(ws) + "."


def paragraph_text(rng: random.Random) -> str:
    return " ".join(sentence(rng, rng.randint(9, 20)) for _ in range(rng.randint(3, 6)))


def download_images(n: int) -> list[bytes]:
    imgs: list[bytes] = []
    for i in range(n):
        seed = IMAGE_SEEDS[i % len(IMAGE_SEEDS)]
        w, h = RNG.choice([(900, 600), (800, 500), (1000, 640), (720, 720)])
        url = f"https://picsum.photos/seed/{seed}/{w}/{h}"
        try:
            with urllib.request.urlopen(url, timeout=30) as resp:
                imgs.append(resp.read())
            print(f"  downloaded {seed} ({w}x{h})", file=sys.stderr)
        except Exception as exc:  # noqa: BLE001
            print(f"  image {seed} failed: {exc}", file=sys.stderr)
    return imgs


def build_model(images: list[bytes]) -> list[dict]:
    """A list of block dicts: heading | para | table | image."""
    rng = random.Random(1)
    blocks: list[dict] = [{"type": "title", "text": "Master Services Agreement"}]
    sections = SECTIONS
    for s in range(sections):
        blocks.append({"type": "heading", "text": f"{s + 1}. {rng.choice(HEADINGS)}"})
        for _ in range(rng.randint(5, 9)):
            blocks.append({"type": "para", "text": paragraph_text(rng)})
        if s % 3 == 0:
            rows = rng.randint(4, 7)
            table = [["Ref", "Description", "Owner", "Value"]]
            for r in range(rows):
                table.append([
                    f"{s + 1}.{r + 1}",
                    sentence(rng, rng.randint(3, 6)),
                    rng.choice(WORDS).capitalize(),
                    f"${rng.randint(1, 999)}k",
                ])
            blocks.append({"type": "table", "rows": table})
        if s % 4 == 1 and images:
            blocks.append({"type": "image", "idx": s % len(images),
                           "caption": f"Figure {s + 1}: {sentence(rng, 4)}"})
    return blocks


def revise_model(blocks: list[dict], images: list[bytes]) -> list[dict]:
    """B: a realistic revision of A's model."""
    rng = random.Random(2)
    out: list[dict] = []
    for b in blocks:
        nb = {**b}
        if b["type"] == "para" and rng.random() < 0.18:
            nb["text"] = paragraph_text(rng)  # reword
        elif b["type"] == "table":
            rows = [list(b["rows"][0])]
            for row in b["rows"][1:]:
                r = list(row)
                if rng.random() < 0.5:
                    r[3] = f"${rng.randint(1, 999)}k"  # changed value
                if rng.random() < 0.3:
                    r[1] = sentence(rng, rng.randint(3, 6))  # changed description
                rows.append(r)
            if rng.random() < 0.4:  # add a row
                rows.append([f"{rows[-1][0]}x", sentence(rng, 4),
                             rng.choice(WORDS).capitalize(), f"${rng.randint(1, 999)}k"])
            nb["rows"] = rows
        elif b["type"] == "image" and images and rng.random() < 0.5:
            nb["idx"] = (b["idx"] + 7) % len(images)  # swapped image
        # occasional insertion/deletion
        if b["type"] == "para" and rng.random() < 0.04:
            continue  # delete
        out.append(nb)
        if b["type"] == "para" and rng.random() < 0.04:
            out.append({"type": "para", "text": paragraph_text(rng)})  # insert
    return out


def render(blocks: list[dict], images: list[bytes], path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(11)
    for b in blocks:
        if b["type"] == "title":
            h = doc.add_heading(b["text"], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif b["type"] == "heading":
            doc.add_heading(b["text"], level=1)
        elif b["type"] == "para":
            doc.add_paragraph(b["text"])
        elif b["type"] == "table":
            rows = b["rows"]
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    t.rows[ri].cells[ci].text = str(cell)
            doc.add_paragraph()
        elif b["type"] == "image":
            data = images[b["idx"]]
            doc.add_picture(io.BytesIO(data), width=Inches(4.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(b.get("caption", ""))
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x73, 0x73, 0x73)
    doc.save(path)
    print(f"wrote {path} ({path.stat().st_size // 1024} KB, {len(blocks)} blocks)", file=sys.stderr)


print("downloading images…", file=sys.stderr)
images = download_images(28)
if not images:
    raise SystemExit("no images downloaded")

model_a = build_model(images)
model_b = revise_model(model_a, images)
render(model_a, images, OUT_DIR / f"sample-a{SUFFIX}.docx")
render(model_b, images, OUT_DIR / f"sample-b{SUFFIX}.docx")
print("done", file=sys.stderr)
