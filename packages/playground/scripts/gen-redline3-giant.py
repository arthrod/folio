"""Generate two *gigantic* (~200-page) themed .docx pairs for the redline3
playground's stress options. Each pair (A + B) is a distinct document, and the
two pairs use different themes so they are genuinely two different giant
examples. B is a realistic revision of A (paragraphs reworded, tables changed,
figures swapped, blocks inserted/deleted). Images are drawn locally with Pillow
so the script never touches the network. Deterministic.

Usage:  python gen-redline3-giant.py <out_dir> [sections]
"""

from __future__ import annotations

import io
import random
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

OUT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(".")
SECTIONS = int(sys.argv[2]) if len(sys.argv) > 2 else 170
OUT_DIR.mkdir(parents=True, exist_ok=True)

BASE_WORDS = (
    "agreement clause party liability warranty covenant termination confidential "
    "obligation governing remedy breach provision counterparty schedule appendix "
    "representation disclosure amendment consideration enforceable severable "
    "assignment notice waiver compliance instrument delivery acceptance milestone "
    "deliverable precedent statutory equitable fiduciary resolution ratify rescind "
    "encumbrance collateral guarantor recital whereas hereto therein hereunder "
    "indemnity jurisdiction arbitration royalty license premises escalation quorum"
).split()

GIANTS = [
    {
        "slug": "giant1",
        "title": "Master Services Agreement (Consolidated)",
        "accent": (37, 99, 235),
        "headings": [
            "Definitions", "Scope of Services", "Service Levels", "Fees and Payment",
            "Deliverables", "Acceptance Testing", "Warranties", "Indemnification",
            "Limitation of Liability", "Confidentiality", "Intellectual Property",
            "Data Protection", "Term and Termination", "Governing Law",
            "Dispute Resolution", "Force Majeure", "Assignment", "Notices",
            "Audit Rights", "Schedules",
        ],
        "extra": "engagement statement-of-work onboarding uptime escalation".split(),
    },
    {
        "slug": "giant2",
        "title": "Senior Secured Credit Facility Agreement",
        "accent": (5, 150, 105),
        "headings": [
            "Definitions and Construction", "The Facilities", "Purpose",
            "Conditions Precedent", "Utilisation", "Repayment", "Prepayment",
            "Interest", "Fees", "Tax Gross-Up", "Representations", "Covenants",
            "Events of Default", "Guarantee and Indemnity", "Security",
            "Changes to Lenders", "Agency", "Set-Off", "Notices", "Governing Law",
        ],
        "extra": "lender borrower facility drawdown margin tranche collateral covenant".split(),
    },
]


def sentence(rng: random.Random, words: list[str], n: int) -> str:
    ws = [rng.choice(words) for _ in range(n)]
    ws[0] = ws[0].capitalize()
    return " ".join(ws) + "."


def paragraph(rng: random.Random, words: list[str]) -> str:
    return " ".join(sentence(rng, words, rng.randint(9, 20)) for _ in range(rng.randint(3, 6)))


def make_image(accent: tuple[int, int, int], label: str) -> bytes:
    w, h = 900, 540
    img = Image.new("RGB", (w, h))
    px = img.load()
    ar, ag, ab = accent
    for y in range(h):
        t = y / h
        for x in range(w):
            s = (x / w) * 0.5 + t * 0.5
            px[x, y] = (
                int(ar * (1 - s) + 245 * s),
                int(ag * (1 - s) + 245 * s),
                int(ab * (1 - s) + 245 * s),
            )
    draw = ImageDraw.Draw(img)
    for i in range(6):
        yy = int(h * (0.12 + i * 0.15))
        draw.line([(0, yy), (w, yy - 70)], fill=(255, 255, 255), width=2)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 64)
    except Exception:  # noqa: BLE001
        font = ImageFont.load_default()
    draw.text((44, h - 108), label, fill=(255, 255, 255), font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def build_model(giant: dict, images: list[bytes]) -> list[dict]:
    rng = random.Random(hash(giant["slug"]) & 0xFFFF)
    words = BASE_WORDS + giant["extra"]
    blocks: list[dict] = [{"type": "title", "text": giant["title"]}]
    for s in range(SECTIONS):
        blocks.append({"type": "heading", "text": f"{s + 1}. {giant['headings'][s % len(giant['headings'])]}"})
        for _ in range(rng.randint(5, 9)):
            blocks.append({"type": "para", "text": paragraph(rng, words)})
        if s % 3 == 0:
            rows = [["Ref", "Description", "Owner", "Value"]]
            for r in range(rng.randint(4, 7)):
                rows.append([
                    f"{s + 1}.{r + 1}",
                    sentence(rng, words, rng.randint(3, 6)),
                    rng.choice(words).capitalize(),
                    f"${rng.randint(1, 999)}k",
                ])
            blocks.append({"type": "table", "rows": rows})
        if s % 4 == 1 and images:
            blocks.append({"type": "image", "idx": s % len(images),
                           "caption": f"Figure {s + 1}: {sentence(rng, words, 4)}"})
    return blocks


def revise_model(giant: dict, blocks: list[dict], images: list[bytes]) -> list[dict]:
    rng = random.Random((hash(giant["slug"]) & 0xFFFF) ^ 0x5A5A)
    words = BASE_WORDS + giant["extra"]
    out: list[dict] = []
    for b in blocks:
        nb = {**b}
        if b["type"] == "para" and rng.random() < 0.18:
            nb["text"] = paragraph(rng, words)
        elif b["type"] == "table":
            rows = [list(b["rows"][0])]
            for row in b["rows"][1:]:
                r = list(row)
                if rng.random() < 0.5:
                    r[3] = f"${rng.randint(1, 999)}k"
                if rng.random() < 0.3:
                    r[1] = sentence(rng, words, rng.randint(3, 6))
                rows.append(r)
            if rng.random() < 0.4:
                rows.append([f"{rows[-1][0]}x", sentence(rng, words, 4),
                             rng.choice(words).capitalize(), f"${rng.randint(1, 999)}k"])
            nb["rows"] = rows
        elif b["type"] == "image" and images and rng.random() < 0.5:
            nb["idx"] = (b["idx"] + 3) % len(images)
        if b["type"] == "para" and rng.random() < 0.04:
            continue
        out.append(nb)
        if b["type"] == "para" and rng.random() < 0.04:
            out.append({"type": "para", "text": paragraph(rng, words)})
    return out


def render(blocks: list[dict], images: list[bytes], path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(11)
    for b in blocks:
        if b["type"] == "title":
            h = doc.add_heading(b["text"], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif b["type"] == "heading":
            doc.add_heading(b["text"], level=1)
        elif b["type"] == "para":
            doc.add_paragraph(b["text"])
        elif b["type"] == "table":
            rows = b["rows"]
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    t.rows[ri].cells[ci].text = str(cell)
            doc.add_paragraph()
        elif b["type"] == "image":
            doc.add_picture(io.BytesIO(images[b["idx"]]), width=Inches(4.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(b.get("caption", ""))
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x73, 0x73, 0x73)
    doc.save(path)
    print(f"wrote {path.name} ({path.stat().st_size // 1024} KB, {len(blocks)} blocks)", file=sys.stderr)


for giant in GIANTS:
    imgs = [make_image(giant["accent"], f"{giant['title'][:24]} · Fig {j + 1}") for j in range(8)]
    model_a = build_model(giant, imgs)
    model_b = revise_model(giant, model_a, imgs)
    render(model_a, imgs, OUT_DIR / f"{giant['slug']}-a.docx")
    render(model_b, imgs, OUT_DIR / f"{giant['slug']}-b.docx")
print("done", file=sys.stderr)
